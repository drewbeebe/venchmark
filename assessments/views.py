#global use imports
from __future__ import print_function
from __future__ import unicode_literals
import os
from io import BytesIO
from subprocess import Popen
from subprocess import PIPE
from subprocess import STDOUT
import base64
import html
import docx
#import django-storages
#from storages.backends.s3boto3 import S3Boto3Storage
#from storages.backends import s3boto3

##imports for secretfileview
import logging
# import boto3
# from botocore.exceptions import ClientError
# import requests

#imports for django
#import datetime as date
from datetime import datetime, date
#from datetime import datetime
from django.shortcuts import render
from django.template import RequestContext
from django.http import HttpResponseRedirect, HttpResponse, FileResponse, StreamingHttpResponse
from django.urls import reverse, reverse_lazy
from django.views.generic import View, DetailView, CreateView, ListView, DeleteView, FormView, UpdateView, TemplateView, RedirectView
from django.views.generic.edit import UpdateView
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin
from rest_framework.authtoken.models import Token
from django.conf import settings
from braces.views import GroupRequiredMixin
from bootstrap_modal_forms.generic import BSModalCreateView, BSModalDeleteView, BSModalFormView
from django.shortcuts import redirect, render
from django.db.models import Q
from django.db.models.signals import post_save
from django.core.paginator import Paginator
from django.core.mail import send_mail
from django.shortcuts import get_object_or_404
from django.templatetags.static import static
from django.core.exceptions import PermissionDenied
from django.views.decorators.csrf import csrf_exempt
#from bootstrap_modal_forms.generic import BSModalView #, BSModalCreateView
from django.forms.models import inlineformset_factory
#from django.core.servers.basehttp import FileWrapper
from wsgiref.util import FileWrapper
from venchmark.settings import BASE_DIR   #, AWS_S3_BUCKET_URL
from django.contrib.staticfiles import finders
from django.core.files.base import ContentFile
from docx.shared import Inches
from docx.enum.style import WD_STYLE
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
#from storages.backends.s3boto3 import S3Boto3Storage

#imports for PDF
from django.utils import timezone
from django.utils.html import format_html
from .render import Render
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.pagesizes import A4
from wkhtmltopdf.views import PDFTemplateView
from .html2docx import html2docx

#from django.shortcuts import render
#from templated_docs import fill_template
#from templated_docs.http import FileResponse

#import cStringIO as StringIO
from io import StringIO, BytesIO
from xhtml2pdf import pisa
from django.template.loader import get_template
from django.template import Context
#from django.http import HttpResponse
#from cgi import escape
from html import escape
from docx.enum.text import WD_ALIGN_PARAGRAPH


#local imports
from .forms import NewAssessmentForm, NewFrameworkForm, NewFrameworkSourceForm, VendorContactQuestionnaireForm, AnalysisQuestionnaireForm, CreateQuestionsForm, NewAnswerForm, NewAnalysisForm, NewAnalysisForm2, UpdateFrameworkForm, UpdateAssessmentForm, CreateReportForm, UpdateFrameworkControlsForm, Add2FWFrameworkControlsForm, AssessmentDetailForm, AssessmentStatusForm, AssessmentDeleteForm, QuestionEditForm, AddMulti2FWFrameworkControlsForm, AnswerUpdateForm, DocumentForm, AnswerDocumentUpdateForm, BSModAnswerDocumentUpdateForm, ReportUpdateForm, ReportGenForm, QuestionEditUpdateForm
from . models import Assessment, Framework, FrameworkSource, FrameworkControls, AssessmentFrameworks, Questionnaire, Question, Report, Document
from companies.models import Company, User
from .utils import handle_uploaded_file, PopulateQuestions
#from django_weasyprint import WeasyTemplateResponseMixin
#from django_weasyprint.views import CONTENT_TYPE_PNG
#from .htmldocxconv import htmldocxconv

#views go here
class AssessmentListView(GroupRequiredMixin, ListView):
    group_required = [u"analyst", u"owner", u"auditor"]
    login_url = '/'
    model = Assessment
    queryset = Assessment.objects.all()
    #fields = ['uuid', 'Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    #fields = ['uuid', 'Name' ]
    fields = ['name', 'vendor', 'owner', 'start_date', 'complete_date', 'status']
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

    def get_queryset(self):
        #myAssessments = Assessment.objects.filter(owner=self.request.user) | Assessment.objects.filter(vendor_contact=self.request.user)
        #return myAssessments
        if self.request.user.is_superuser:
            myAssessments = Assessment.objects.all()
        else:
            print(str(self.request.user.last_name))
            print(str(self.request.user.groups.all))
            myAssessments = Assessment.objects.filter(Q(vendor_contact=self.request.user) | Q(analyst=self.request.user) | Q(owner=self.request.user))
        return myAssessments


class AssessmentFormActionMixin(object):

    def post(self, request, *args, **kwargs):
        """Add 'Cancel' button redirect."""
        if "cancel" in request.POST:
            url = reverse('assessment_list')     # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
        else:
            return super(AssessmentFormActionMixin, self).post(request, *args, **kwargs)

class FrameworkControlAddFormActionMixin(object):

    def post(self, request, *args, **kwargs):
        """Add 'Cancel' button redirect."""
        if "cancel" in request.POST:
            #slug_url_kwarg = FrameworkControls.uuid
            #control = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
            #frameworkuuid = control.framework.uuid
            url = '/controls/' + str(self.kwargs['uuid'])    # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
            #url = reverse('framework_list')     # or e.g. reverse(self.get_success_url())
            #return HttpResponseRedirCannot assign "'MyCo Documented Framework'": "FrameworkControls.framework" must be a "Framework" instanect(url)
        else:
            return super(FrameworkControlAddFormActionMixin, self).post(request, *args, **kwargs)


class FrameworkControlUpdateFormActionMixin(object):

    def post(self, request, *args, **kwargs):
        """Add 'Cancel' button redirect."""
        if "cancel" in request.POST:
            slug_url_kwarg = FrameworkControls.uuid
            control = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
            #frameworkuuid = control.framework.uuid
            url = '/controls/' + str(control.frameworkUUID)    # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
            #url = reverse('framework_list')     # or e.g. reverse(self.get_success_url())
            #return HttpResponseRedirect(url)
        else:
            return super(FrameworkControlUpdateFormActionMixin, self).post(request, *args, **kwargs)

class FrameworkFormActionMixin(object):

    def post(self, request, *args, **kwargs):

###### THIS CODE HERE IS FUCKED AND DOESN"T WORK....FIND CASE LOGIC AND APPLY IT
        """Add 'AddSource' button redirect."""
        if "AddSource" in request.POST:
            url = reverse('frameworksource_new')     # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
        elif "cancel" in request.POST:
            url = reverse('framework_list')     # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
        elif "AddControls" in request.POST:
            slug_url_kwarg = Framework.uuid
            url = '/controls/' + str(self.kwargs['uuid']) + '/add/'    # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
        else:
            return super(FrameworkFormActionMixin, self).post(request, *args, **kwargs)



class FrameworkSourceFormActionMixin(object):

    def post(self, request, *args, **kwargs):
        """Add 'Cancel' button redirect."""
        if "cancel" in request.POST:
            url = reverse('framework_list')     # or e.g. reverse(self.get_success_url())
            return HttpResponseRedirect(url)
        else:
            return super(FrameworkSourceFormActionMixin, self).post(request, *args, **kwargs)


class AssessmentCreateView(GroupRequiredMixin, AssessmentFormActionMixin, CreateView):
    group_required = [ u"owner", u"administrator" ]
    login_url = '/'
    model = Assessment
    #template_name = 'Framework_new.html'
    #template_name = 'assessments/NewAssessmentForm.htFrameworkFormActionMixinml'
    template_name = 'assessments/assessment_create_form.html'
    form_class = NewAssessmentForm
    #fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Frameworks']



    def get_success_url(self):
            return reverse('assessment_list')

    def get_initial(self):
        #return {'Owner': self.kwargs['request.user']}
        return {'owner': self.request.user,
                'auditor': self.request.user,
                'analyst': self.request.user}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        vendors = Company.objects.all()
        context['vendors'] = vendors
        return context
    #def populate_questionnaire():
    #    newQuestionnaire = Questionnaire()


    #post_save.connect(populate_questionnaire, sender=Assessment)



class AssessmentDeleteView(GroupRequiredMixin, DeleteView):
    group_required = [ u"owner", u"administrator" ]
    login_url = '/'
    model = Assessment
    slug_url_kwarg = Assessment.uuid

    def get_success_url(self):
            return reverse('assessment_list')

    def get_object(self, queryset=None):
        obj = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Assessment.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

#### 2020-04-05 - DB - Migrating Frameworks views to Assessments Views


class AssessmentStatusView(GroupRequiredMixin, FormView):
    group_required = [ u"owner", u"auditor" ]
    login_url = '/'
    redirect_unauthenticated_users = True
    model = Assessment
    slug_url_kwarg = Assessment.uuid
    form_class = AssessmentStatusForm
    template_name = "assessments/assessment_status.html"

    #def get_object(self, queryset=None):
    #    obj = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
    #    return obj

    def get_context_data(self, **kwargs):
        context = super(AssessmentStatusView, self).get_context_data(**kwargs)
        This_questionnaire = Questionnaire.objects.filter(assessment=self.kwargs['uuid']).first()
        context['questionnaire'] = This_questionnaire
        This_assessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        context['assessment'] = This_assessment

        return context


class FrameworkListView(GroupRequiredMixin, ListView):
    group_required = [ u"administrator" ]
    #authenticated_redirect_url = u"/"
    #login_url = '/users/login/'
    login_url = '/'
    redirect_unauthenticated_users = True
    model = Framework
    #context_object_name = 'Framework_list'   # your own name for the list as a template variable
    queryset = Framework.objects.all()
    fields = ['Name', 'ShortName', 'Version', 'PublishDate', 'Source']
    #template_name = 'Frameworks/Frameworks_list.html'  # Specify your own template name/location
    paginate_by = 10

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context


class FrameworkDetailView(GroupRequiredMixin, DetailView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = Framework
    slug_url_kwarg = Framework.uuid
    fields = ['name', 'short_name', 'version', 'publish_date', 'source']
    def get_object(self, queryset=None):
        obj = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Framework.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        print("frameworkUUID = " + str(self.kwargs['uuid']))
        #controls = self.object.frameworkcontrols_set.filter(frameworkUUID=self.kwargs['uuid']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context




class FrameworkUpdateView(GroupRequiredMixin, FrameworkFormActionMixin, UpdateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = Framework
    slug_url_kwarg = Framework.uuid
    template_name = 'assessments/framework_update.html'
    fields = '__all__'
    def get_success_url(self):
            return reverse('framework_detail')
    def get_object(self, queryset=None):
        obj = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        #print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Framework.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #print("frameworkUUID = " + str(self.kwargs['uuid']))
        #controls = self.object.frameworkcontrols_set.filter(frameworkUUID=self.kwargs['uuid']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class FrameworkUpdateView2(GroupRequiredMixin, FrameworkFormActionMixin, UpdateView):
#class PersonUpdateView(BSModalUpdateView):
    group_required = [ u"owner", u"administrator"]
    login_url = '/'
    model = Framework
    slug_url_kwarg = Framework.uuid
    template_name = 'assessments/framework_update_form.html'
    form_class = UpdateFrameworkForm
    success_url = reverse_lazy('framework_list')
    #fields = '__all__'

    #def get_success_url(self):
    #        return reverse('user_detail')

    def get_object(self, queryset=None):
        #obj = User.objects.filter(uuid=self.kwargs['uuid']).first()
        myFramework = Framework.objects.get(pk=self.kwargs['uuid'])
        #print(myUser.first_name + " " + myUser.last_name)
        #groups = myUser.groups.all()
        #for group in groups:
        #    print(group)
        #print(groups)
        #print("user belongs " + str(groups.count()) + " groups.")
        #for group in groups:
    #        print("retrieved group" + str(group))
        return myFramework

#    def form_valid(self, form):
        #self.object.groups.clear()
        #groups = form.cleaned_data['groups']
        #for group in groups:
        #    print("saving this group in user: " + str(group))
        #    user.object.groups.add(group)
#        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
#        return super().form_valid(form)

class FrameworkDeleteView(GroupRequiredMixin, DeleteView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = Framework
    slug_url_kwarg = Framework.uuid
    #template_name = 'Framework_delete.html'
    def get_success_url(self):
            return reverse('framework_list')

    def get_object(self, queryset=None):
        obj = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Framework.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        print("frameworkUUID = " + str(self.kwargs['uuid']))
        #controls = self.object.frameworkcontrols_set.filter(frameworkUUID=self.kwargs['uuid']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class FrameworkSourceListView(GroupRequiredMixin, ListView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkSource

    #context_object_name = 'FrameworkSource_list'   # your own name for the list as a template variable
    queryset = FrameworkSource.objects.all()
    fields = ['Name', 'Acronym', 'URL']
    #template_name = 'FrameworkSources/FrameworkSources_list.html'  # Specify your own template name/location


class FrameworkSourceDetailView(GroupRequiredMixin, DetailView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkSource
    slug_url_kwarg = FrameworkSource.uuid
    #template_name = 'FrameworkSource_detail.html'
    fields = ['Name', 'Acronym', 'URL']
    def get_object(self, queryset=None):
        obj = FrameworkSource.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        queryset = FrameworkSource.objects.filter(uuid=self.kwargs['uuid']).first()
        return queryset


class FrameworkSourceCreateView(GroupRequiredMixin, CreateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkSource
    #template_name = 'FrameworkSource_new.html'
    fields = ['Name', 'Acronym', 'URL']
    def get_success_url(self):
            #return reverse('frameworksource_list')
            return reverse('framework_list')

class FrameworkSourceUpdateView(GroupRequiredMixin, UpdateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkSource
    template_name = 'frameworksource_update.html'
    fields = '__all__'
    def get_success_url(self):
            return reverse('frameworksource_detail')

class FrameworkSourceDeleteView(GroupRequiredMixin, DeleteView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkSource

    #template_name = 'FrameworkSource_delete.html'
    def get_success_url(self):
            return reverse('frameworksource_list')

class FrameworkControlsListView(GroupRequiredMixin, ListView):
    group_required = [ u"administrator" ]
    login_url = '/'
    #slug_url_kwarg = FrameworkControls.frameworkUUID
    slug_url_kwarg = FrameworkControls.frameworkUUID
    model = FrameworkControls
    template_name = 'assessments/controls_list.html'
    context_object_name = 'frameworkcontrols'  # Default: object_list
    #context_object_name = 'FrameworkControls_list'   # your own name for the list as a template variable
    #queryset = FrameworkControls.objects.filter(FrameworkControlsID=uuid)
    fields = ['FrameworkControlsID', 'Function', 'Category', 'SubCategoryID', 'SubCategory', 'Reference', 'DefaultQuestion']
    #template_name = 'FrameworkControlss/FrameworkControlss_list.html'  # Specify your own template name/location
    #paginate_by = 10
    paginate_by = 10

    def get_queryset(self):
        print(str(self.kwargs['uuid']))
        queryset = FrameworkControls.objects.filter(frameworkUUID=self.kwargs['uuid'])
        print("Items in queryset: " + str(queryset.count()))
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #print("frameworkUUID = " + str(self.kwargs['uuid']))
        thisframework = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        context['framework'] = thisframework
        return context


class FrameworkControlsDetailView(GroupRequiredMixin, DetailView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    #template_name = 'FrameworkControls_detail.html'
    fields = ['framework', 'Function', 'Category', 'SubCategoryID', 'SubCategory', 'Reference', 'DefaultQuestion']
    slug_url_kwarg = FrameworkControls.uuid
    def get_object(self, queryset=None):
        obj = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class FrameworkControlsCreateView( GroupRequiredMixin, CreateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    #template_name = 'FrameworkControls_new.html'
    fields = [ 'framework', 'function', 'functionID', 'category', 'categoryID', 'category_statement', 'subcategory', 'subcategoryID', 'control_statement', 'default_question', 'reference' ]
    def get_success_url(self):
            return reverse('FrameworkControls_list')

class FrameworkControlsAddView( GroupRequiredMixin, FrameworkControlAddFormActionMixin, CreateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    slug_url_kwarg = Framework.uuid
    template_name = 'assessments/frameworkcontrols_add.html'
    form_class = Add2FWFrameworkControlsForm
    #template_name = 'FrameworkControls_new.html'
#    fields = [ 'framework', 'function', 'functionID', 'category', 'categoryID', 'category_statement', 'subcategory', 'subcategoryID', 'control_statement', 'default_question', 'reference' ]



    def get_success_url(self):
        success_url = "/controls/" + str(self.kwargs['uuid']) + "/"
        return success_url #reverse('FrameworkControls_list')

    def get_initial(self):
        framework = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        frameworkUUID = self.kwargs['uuid']
    #    #print(framework.name)
        return {'framework': framework, 'frameworkUUID': frameworkUUID}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        framework = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        context['framework'] = framework
        return context

def upload_multiple_controls_file(request, uuid):
    context = {}

    # fetch the object related to passed id
    obj = get_object_or_404(Framework, uuid = uuid)

    template_name = 'assessments/frameworkcontrols_multiadd.html'
    #framework = Framework.objects.filter(uuid=request.kwargs['uuid']).first()
    framework = Framework.objects.filter(uuid=uuid).first()
    #print(framework.name)
    #form = AddMulti2FWFrameworkControlsForm(request.POST or None, instance = obj)
    initial_dict = {
        "framework" : framework.uuid
    }


    if request.method == 'POST':
        print("method is post")
        framework = Framework.objects.filter(uuid=uuid).first()
        FrameworkControls.framework = uuid
        form = AddMulti2FWFrameworkControlsForm(request.POST, request.FILES, initial = initial_dict)

        if form.is_valid():
            print("form is valid")
            #"uploads/{}/{}".format(instance.uuid, filename)
            #form.save()

            handle_uploaded_file(request.FILES['controls_file'], uuid)
            return HttpResponseRedirect('/controls/' + str(uuid) + '/')
    else:
        form = AddMulti2FWFrameworkControlsForm(initial = initial_dict)

    context['form'] = form
    context['framework'] = obj
    return render(request, template_name, context)


class FrameworkControlsAddMultipleView( GroupRequiredMixin, FrameworkControlAddFormActionMixin, CreateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    slug_url_kwarg = Framework.uuid
    template_name = 'assessments/frameworkcontrols_multiadd.html'
    form_class = AddMulti2FWFrameworkControlsForm
    #framework = Framework.objects.get(uuid=request.POST["uuid"])
    #template_name = 'FrameworkControls_new.html'
#    fields = [ 'framework', 'function', 'functionID', 'category', 'categoryID', 'category_statement', 'subcategory', 'subcategoryID', 'control_statement', 'default_question', 'reference' ]
    def form_valid(self, form):
        # This method is called when valid form data has been POSTed.
        # It should return an HttpResponse.
        framework = Framework.objects.get(uuid=request.POST["uuid"])
        handle_uploaded_file(request.FILES['controls_file'], self.kwargs['uuid'])
        return super().form_valid(form)

    def get_success_url(self):
            return reverse('FrameworkControls_list')

    def get_initial(self):
        framework = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        #print(framework.name)
        return {'framework': framework}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        framework = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        context['framework'] = framework
        return context


class FrameworkControlsUpdateView(GroupRequiredMixin, FrameworkControlUpdateFormActionMixin, UpdateView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    template_name = 'assessments/frameworkcontrols_update.html'
    slug_url_kwarg = FrameworkControls.uuid
    #fields = '__all__'
    form_class = UpdateFrameworkControlsForm
    def get_success_url(self):
            #return reverse('frameworkControls_detail')
            urlstring = "/controls/" + str(self.kwargs['uuid'])
            return HttpResponse(urlstring)

    def get_object(self, queryset=None):
        obj = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['id'])
        queryset = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return queryset

class FrameworkControlsDeleteView(GroupRequiredMixin, DeleteView): #DeleteView):
    group_required = [ u"administrator" ]
    login_url = '/'
    model = FrameworkControls
    template_name = 'assessments/controls_delete_form.html'

    #success_url = "/controls/" + str(request.kwargs['uuid']) + "/" #reverse_lazy('controls_list')
    #def get_success_url(self):
        #success_url = "/controls/" + str(self.request.kwargs['uuid']) + "/"
        #context = self.get_context_data()
        #framework = context['framework']
        #print("context framework = " + str(framework.name))
        #print("uuid of context framework = " + str(framework.uuid))
        #success_url = "/controls/" + str(framework.uuid) + "/"
    #    success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
    #    return success_url

    def get_object(self, queryset=None):
        obj = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj


    def get_queryset(self):
        #print("self kwarg id:" + self.kwargs['id'])
        queryset = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        control = FrameworkControls.objects.filter(uuid=self.kwargs['uuid']).first()
        framework = control.framework
        context['control'] = control
        context['framework'] = framework
        return context

    def get_success_url(self):
        uuid = self.kwargs['uuid']
        return HttpResponseRedirect('/controls/' + str(uuid))


class FrameworkCreateView(GroupRequiredMixin, FrameworkFormActionMixin, CreateView):
    group_required = [ u"administrator" ]
    template_name = 'assessments/NewFrameworkForm.html'
    form_class = NewFrameworkForm
    success_message = 'Success: Framework was created.'
    success_url = reverse_lazy('framework_list')

class FrameworkSourceCreateView(GroupRequiredMixin, FrameworkSourceFormActionMixin, CreateView):
    group_required = [ u"administrator" ]
    template_name = 'assessments/NewFrameworkSourceForm.html'
    form_class = NewFrameworkSourceForm
    success_message = 'Success: Framework Source was created.'
    #success_url = reverse_lazy('source_list')
    success_url = reverse_lazy('framework_list')



class QuestionsCreateView(GroupRequiredMixin, FormView):
    group_required = [ u"owner", u"administrator" ]
    login_url = '/'
    #success_url = '/assessment/'+str(self.kwargs['uuid'])+'/'
    #template_name = 'Framework_new.html'
    template_name = 'assessments/NewQuestionsForm.html'
    form_class = CreateQuestionsForm

    def form_valid(self, form):
        # This method is called when valid form data has been POSTed.
        # It should return an HttpResponse.
        form.PopulateQuestions(self.kwargs['uuid'])
        return super().form_valid(form)
    #fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Frameworks']
    def get_success_url(self):
            return '/assessment/'+str(self.kwargs['uuid'])+'/'   #reverse('assessment_detail'+'/'+str(self.kwargs['uuid'])+'/')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #print("frameworkUUID = " + str(self.kwargs['uuid']))

        qassessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        questionnaire = Questionnaire.objects.filter(assessment=qassessment)
        frameworks = qassessment.frameworks.all()
        context['questionnaire'] = questionnaire
        context['assessment'] = qassessment
        context['frameworks'] = frameworks
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context


#Questionnaire Views
## VendorContactView -
class VendorQuestionnaireView(GroupRequiredMixin, DetailView):
    group_required = [ u"owner", u"vendor", u"administrator" ]
    login_url = '/'

    template_name = 'assessments/questions_detail.html'
    form_class = VendorContactQuestionnaireForm

    slug_url_kwarg = Question.questionnaire


    def get_object(self, queryset=None):
        obj = Question.objects.filter(questionnaire_id=self.kwargs['uuid']).all()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Question.objects.filter(questionnaire_id=self.kwargs['uuid']).all()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        questions = Question.objects.filter(questionnaire_id=self.kwargs['uuid'])
        context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class VendorQuestionListView(GroupRequiredMixin, ListView):
    group_required = [ u"administrator", u"vendor", u"owner", u"analyst" ]
    #authenticated_redirect_url = u"/"
    #login_url = '/users/login/'
    login_url = '/'
    redirect_unauthenticated_users = True
    slug_url_kwarg = Assessment.uuid
    model = Question
    paginate_by = 10



    def get_template_names(self):
        user=self.request.user
        assessment=Assessment.objects.filter(uuid=self.kwargs['uuid']).first()

        if user.groups.filter(name=u"owner").exists():
            print("User is in owner group")
        if user.groups.filter(name=u"vendor").exists():
            print("User is in vendor group")
        if user.is_superuser:
            print("User is superuser")

        print("assessment status: " + str(assessment.status))

        if user.groups.filter(name = u"owner").exists() and assessment.status == "QUESTIONNAIRE_REVIEW":
            #print("both conditions met - setting template name")
            template_name = "assessments/questions_list_owner_q_review.html"
        elif user.groups.filter(name = u"owner").exists() and assessment.status == "VENDOR_SUBMIT":
            #print("both conditions met - setting template name")
            template_name = "assessments/questions_list_owner_vendor_submit.html"
        elif user.groups.filter(name = u"owner").exists():
            #print("both conditions met - setting template name")
            template_name = "assessments/questions_list_owner.html"

        elif user.groups.filter(name = u"vendor").exists() and assessment.status == "VENDOR_SUBMIT":
            #print("both conditions met - setting template name")
            template_name = "assessments/questions_list_vendor_vendor_submit.html"
        elif user.is_superuser and assessment.status == "IN_ANALYSIS":
            template_name = "assessments/questions_list_analyst_analysis.html"
        elif user.groups.filter(name = u"analyst").exists() and assessment.status == "IN_ANALYSIS":
            template_name = "assessments/questions_list_analyst_analysis.html"
        elif user.groups.filter(name = u"vendor").exists() and assessment.status == "IN_ANALYSIS":
            template_name = "assessments/questions_list_vendor_analysis.html"

        elif user.groups.filter(name = u"analyst").exists() and assessment.status == "REPORT_GENERATED":
            report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
            if report:
                template_name = "assessments/questions_list_analyst_report_exists.html"
            else:
                template_name = "assessments/questions_list_analyst_report_doesnt_exist.html"
        #elif user.groups.filter(name = u"analys").exists():
        #    template_name = "assessments/questions_list_analyst.html"
        return template_name
    #form_class = VendorContactQuestionnaireForm
    #elif current_user.groups.filter(name = u"analyst").exists():
    #    form_class = AnalystContactQuestionnaireForm
    #elif django_user.groups.filter(name = u"adminstrator").exists():
    #    form_class = AdministratorContactQuestionnaireForm



    def get_queryset(self):
        #print("self kwarg id:" + self.kwargs['uuid'])
        #queryset = Question.objects.filter(assessment_id=self.kwargs['uuid']).all()
        queryset = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #questions = Question.objects.filter(assessment_id=self.kwargs['uuid']).all()
        questions = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        total_answers = 0
        analysis_count = 0
        analyzed_questions = []
        for question in questions:
            if question.answer:
                total_answers += 1

        for analyzed_question in questions:
            if analyzed_question.was_analyzed:
                analyzed_questions.append(analyzed_question)
                analysis_count += 1

        question_count = len(questions)
        analysis_count = len(analyzed_questions)
        print(analysis_count)
        complete_pct = int((total_answers / question_count) * 100)
        if analysis_count == 0:
            analysis_complete_pct = 0
        else:
            analysis_complete_pct = int(( analysis_count / question_count) * 100)
        context['questions'] = questions
        context['question_count'] = question_count
        context['answer_count'] = total_answers
        context['complete_pct'] = complete_pct
        context['uuid'] = self.kwargs['uuid']
        context['analysis_count'] = analysis_count
        context['analysis_pct'] = analysis_complete_pct
        #token = Token.objects.get(user=self.request.user)
        #context['token'] = token
        #current_page = context.pop('page_obj', None)
        #context['current_page'] = current_page
        return context


class VendorQuestionListView2(GroupRequiredMixin, ListView):
    group_required = [ u"administrator", u"vendor", u"owner", u"analyst" ]
    login_url = '/'
    redirect_unauthenticated_users = True
    slug_url_kwarg = Assessment.uuid
    uuid = slug_url_kwarg
    model = Question
    paginate_by = 10


    def get_template_names(self):
        user=self.request.user
        #assessment=Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        template_name = ""

        class QuestionnaireStatusSwitch:

            def switch(self, status):
                #default = "No Status to Change"
                template_name = "assessments/questionnaire_default.html"
                return getattr(self, status, lambda: default)()

            def CREATED(self):
                template_name = "assessments/questionnaire_created.html"
                return template_name

            def QUESTIONNAIRE_REVIEW(self):
                template_name = "assessments/questionnaire_review.html"
                return template_name

            def VENDOR_SUBMIT(self):
                template_name = "assessments/questionnaire_vendor_submit.html"
                return template_name

            def IN_ANALYSIS(self):
                template_name = "assessments/questionnaire_in_analysis.html"
                return template_name

            def REPORT_GENERATED(self):
                template_name = "assessments/questionnaire_report_generated.html"
                return template_name

            def ASSESSMENT_COMPLETE(self):
                template_name = "assessments/questionnaire_assessment_complete.html"
                return template_name


        thisassessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()

        s = QuestionnaireStatusSwitch()

        template_name = s.switch(thisassessment.status)
        print("template name: " + str(template_name))
        return template_name


    def get_queryset(self):
        #print("self kwarg id:" + self.kwargs['uuid'])
        #queryset = Question.objects.filter(assessment_id=self.kwargs['uuid']).all()
        queryset = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        return queryset

    def get_context_data(self, **kwargs):
        def safe_complete_pct(tot_answ, q_count):
            if tot_answ <= 0:
                return 0
            else:
                compl_pct = int((tot_answ / q_count) * 100)
            return compl_pct
        def safe_analysis_complete(anl_count, q_count):
            if anl_count <= 0:
                return 0
            else:
                anl_count_pct = int(( anl_count / q_count) * 100)
            return anl_count_pct
        context = super().get_context_data(**kwargs)
        #questions = Question.objects.filter(assessment_id=self.kwargs['uuid']).all()
        questions = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        total_answers = 0
        analysis_count = 0
        analyzed_questions = []
        for question in questions:
            if question.answer:
                total_answers += 1

        for analyzed_question in questions:
            if analyzed_question.was_analyzed:
                analyzed_questions.append(analyzed_question)
                analysis_count += 1

        question_count = len(questions)
        analysis_count = len(analyzed_questions)
        #print(analysis_count)
        complete_pct = safe_complete_pct(total_answers, question_count)  #int((total_answers / question_count) * 100)
        analysis_complete_pct = safe_analysis_complete(analysis_count, question_count)

        #if analysis_count == 0:
        #    analysis_complete_pct = 0
        #else:
        #    analysis_complete_pct = int(( analysis_count / question_count) * 100)
        thisassessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        context['assessment'] = thisassessment
        context['questions'] = questions
        context['question_count'] = question_count
        context['answer_count'] = total_answers
        context['complete_pct'] = complete_pct
        context['uuid'] = self.kwargs['uuid']
        context['analysis_count'] = analysis_count
        context['analysis_pct'] = analysis_complete_pct

        return context


class AnalysisQuestionListView(GroupRequiredMixin, ListView):
    group_required = [ u"administrator", u"owner", u"analyst" ]
    #authenticated_redirect_url = u"/"
    #login_url = '/users/login/'
    login_url = '/'
    redirect_unauthenticated_users = True
    slug_url_kwarg = Assessment.uuid
    model = Question
    #if current_user.groups.filter(name = u"vendor").exists():
    form_class = AnalysisQuestionnaireForm
    #elif current_user.groups.filter(name = u"analyst").exists():
    #    form_class = AnalystContactQuestionnaireForm
    #elif django_user.groups.filter(name = u"adminstrator").exists():
    #    form_class = AdministratorContactQuestionnaireForm
    template_name = "assessments/analysis_q_list.html"
    #context_object_name = 'Framework_list'   # your own name for the list as a template variable
    #queryset = Question.objects.filter(questionnaire_id=self.kwargs['uuid'])
    fields = "__all__"
    #fields = ['Name', 'ShortName', 'Version', 'PublishDate', 'Source']
    #template_name = 'Frameworks/Frameworks_list.html'  # Specify your own template name/location
    paginate_by = 10

    def get_queryset(self):

        queryset = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        questions = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        context['questions'] = questions
        context['uuid'] = self.kwargs['uuid']
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        report = Report.objects.filter(assessment=self.kwargs['uuid']).first()
        context['report'] = report
        return context

#class FrameworkCreateView(BSModalCreateView):
#    template_name = 'assessments/NewFrameworkForm.html'
#    form_class = NewFrameworkForm
#    success_message = 'Success: Framework was created.'
#    success_url = reverse_lazy('framework_list')


class AnswerUpdateView(BSModalCreateView):
    template_name = 'assessments/NewAnswerForm.html'
    form_class = NewAnswerForm
    success_message = 'Success: Question was answered.'
    success_url = reverse_lazy('question_list')
    #group_required = [ u"owner", u"administrator" ]
    #login_url = '/'
    #model = Question
    #slug_url_kwarg = Question.uuid
    #template_name = 'question_update.html'
    #fields = [ 'question', 'answer' ]

    #def get_success_url(self):
    #        return reverse('question_list')

    #def get_object(self, queryset=None):
    #    obj = Question.objects.filter(uuid=self.kwargs['uuid']).first()
    #    return obj

    #def get_queryset(self):
    #    print("self kwarg id:" + self.kwargs['uuid'])
    #    queryset = Question.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
    #    return queryset

#class QuestionUpdateView(GroupRequiredMixin, UpdateView):
#    group_required = [ u"owner", u"administrator" ]
#    login_url = '/'
#    slug_url_kwarg = Question.uuid
#    template_name = 'assessments/NewAnswerForm-org.html'
#    form_class = NewAnswerForm
#    model = Question

#    def get_success_url(self):
#            return reverse('questions_list')

#    def get_queryset(self):
#        queryset = Question.objects.filter(uuid=self.kwargs['uuid']).first()
#        return queryset

#    def get_context_data(self, **kwargs):
#        #print('the UUID passed is: ' + str(self.kwargs['uuid']))
#        context = super().get_context_data(**kwargs)
#        question = Question.objects.filter(uuid=self.kwargs['uuid']).first()
#        print("QuestionUpdateView was passed the UUID: " + str(question.uuid))
#        context['question'] = question
#        context['uuid'] = self.kwargs['uuid']
#        token = Token.objects.get(user=self.request.user)
#        context['token'] = token
#        return context


class QuestionUpdateView(GroupRequiredMixin, UpdateView):
    group_required = [ u"owner", u"administrator" ]
    login_url = '/'
    template_name = "assessments/questions_form.html"
    #form_class = NewAnswerForm
    form_class = QuestionEditForm
    slug_url_kwarg = Question.uuid

    #passed_quuid =

    #def post(self, request, uuid):
    #    next = request.POST.get('next', '/')
    #    pass
    #    return HttpResponseRedirect(next)

    #def form_valid(self, form):
        #form.instance.created_by = self.request.user
        #uuid = self.kwargs['quuid']
    #    return HttpResponseRedirect('/questions/' + str(quuid) + "/")#super().form_valid(form)
    def form_valid(self, form):
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        return super().form_valid(form)

    def get_object(self, *args, **kwargs):
        question = Question.objects.get(uuid=self.kwargs['uuid'])
        return question

    #def get_success_url(self):
        #self.success_url = self.request.POST.get('previous_page')
        #return HttpResponseRedirect(str(self.request.path))   #("/questionnaire/" + self.kwargs['next'] + "/")

    def get_context_data(self, **kwargs):
        context = super(QuestionUpdateView, self).get_context_data(**kwargs)
        retrieved_question = self.get_object()
        context['question'] = retrieved_question
        qassessmentid = retrieved_question.assessment.uuid
        context['quuid'] = qassessmentid
        return context

def QuestionUpdateView2(request, uuid):
    #template_name = "assessments/questions_form.html"
    #form_class = NewAnswerForm
    #form_class = QuestionEditForm
    slug_url_kwarg = Question.uuid
    question = Question.objects.get(pk=self.kwargs['uuid'])
    new_answer = data['answer']
    success = question.new_answer(new_answer)
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
    #question.save()



class AnalysisUpdateView(GroupRequiredMixin, UpdateView):
#class AnalysisUpdateView(GroupRequiredMixin, BSModalFormView):
    group_required = [ u"owner", u"administrator", u"analyst" ]
    login_url = '/'
    template_name = "assessments/analysis_form.html"
    #form_class = NewAnswerForm
    #model = Question
    form_class = NewAnalysisForm2
    slug_url_kwarg = Question.uuid

    #passed_quuid =

    #def post(self, request, uuid):
    #    next = request.POST.get('next', '/')
    #    pass
    #    return HttpResponseRedirect(next)

    #def form_valid(self, form):
        #form.instance.created_by = self.request.user
        #uuid = self.kwargs['quuid']
    #    return HttpResponseRedirect('/questions/' + str(quuid) + "/")#super().form_valid(form)
    def form_valid(self, form):
        print("form was VALID")
        #self.kwargs['was_analyzed'] = True
        print(form.cleaned_data['was_analyzed'])
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        return super().form_valid(form)

    def get_object(self, *args, **kwargs):
        question = Question.objects.get(pk=self.kwargs['uuid'])
        return question

    #def get_success_url(self):
        #self.success_url = self.request.POST.get('previous_page')
        #return HttpResponseRedirect(str(self.request.path))   #("/questionnaire/" + self.kwargs['next'] + "/")

    def get_context_data(self, **kwargs):
        context = super(AnalysisUpdateView, self).get_context_data(**kwargs)
        retrieved_question = self.get_object()
        context['question'] = retrieved_question
        qassessmentid = retrieved_question.assessment.uuid

        context['quuid'] = qassessmentid
        this_assessment = Assessment.objects.filter(uuid=qassessmentid).first()

        context['assessment_name'] = this_assessment.name
        return context  #, kwargs

class AnalysisUpdateEvidenceView(GroupRequiredMixin, UpdateView):
#class AnalysisUpdateEvidenceView(GroupRequiredMixin, BSModalFormView):
    group_required = [ u"owner", u"administrator", u"analyst" ]
    login_url = '/'
    model = Question
    template_name = "assessments/question_evidence_form.html"
    #form_class = NewAnswerForm
    #model = Question
    #form_class = NewAnalysisForm2
    fields = [ 'document' ]
    slug_url_kwarg = Question.uuid

    def form_valid(self, form):
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        return super().form_valid(form)

    def get_object(self, *args, **kwargs):
        question = Question.objects.get(uuid=self.kwargs['uuid'])
        return question

    #def get_success_url(self):
        #self.success_url = self.request.POST.get('previous_page')
        #return HttpResponseRedirect(str(self.request.path))   #("/questionnaire/" + self.kwargs['next'] + "/")

    def get_context_data(self, **kwargs):
        if 'view' not in kwargs:
            kwargs['view'] = self
        context = super(AnalysisUpdateEvidenceView, self).get_context_data(**kwargs)
        #retrieved_question = self.get_object()
        #context['question'] = retrieved_question
        #qassessmentid = retrieved_question.assessment.uuid
        #context['quuid'] = qassessmentid
        #this_assessment = Assessment.objects.filter(uuid=qassessmentid).first()
        return context  #, kwargs



class FrameworkDeleteView2(GroupRequiredMixin, BSModalDeleteView):
    group_required = [ u"owner", u"administrator"]
    login_url = '/'
    model = Framework
    slug_url_kwarg = Framework.uuid
    template_name = 'assessments/framework_delete_form.html'
    success_message = 'Framework Deleted Successfully.'
    success_url = '/frameworks/'

    #def get_success_url(self):
    #        return reverse('user_list')

    def get_object(self, queryset=None):
        obj = Framework.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
    #    print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Framework.objects.filter(uuid=self.kwargs['uuid']).first() #.order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        if 'view' not in kwargs:
            kwargs['view'] = self
        return kwargs



class AssessmentDeleteView2(GroupRequiredMixin, DeleteView):
    group_required = [ u"owner", u"administrator"]
    login_url = '/'
    model = Assessment
    #slug_url_kwarg = Assessment.uuid
    #template_name = 'assessments/assessment_delete_form.html'
    #success_message = 'Assessment Deleted Successfully.'
    success_url = '/assessments/'
    #form_class = AssessmentDeleteForm

    def get_object(self, queryset=None):
        obj = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj




class AssessmentUpdateView2(GroupRequiredMixin, AssessmentFormActionMixin, UpdateView):
#class PersonUpdateView(BSModalUpdateView):
    group_required = [ u"owner", u"administrator"]
    login_url = '/'
    model = Assessment
    slug_url_kwarg = Assessment.uuid
    template_name = 'assessments/assessment_update_form.html'
    form_class = UpdateAssessmentForm
    #fields = '__all__'

    def get_success_url(self):
            return reverse('assessment_list')

    def get_object(self, queryset=None):
        #obj = User.objects.filter(uuid=self.kwargs['uuid']).first()
        myAssessment = Assessment.objects.get(pk=self.kwargs['uuid'])
        #print(myUser.first_name + " " + myUser.last_name)
        #groups = myUser.groups.all()
        #for group in groups:
        #    print(group)
        #print(groups)
        #print("user belongs " + str(groups.count()) + " groups.")
        #for group in groups:
    #        print("retrieved group" + str(group))
        return myAssessment

    def form_valid(self, form,  **kwargs):
        context = self.get_context_data(**kwargs)
        initial_vendor_POC = context['initial_vendor_contact']
        init_analyst       = context['initial_analyst']
        thisassessment     = Assessment.objects.filter(name=form.cleaned_data['name']).first()
        #self.object.groups.clear()
        #groups = form.cleaned_data['groups']
        #for group in groups:
        #    print("saving this group in user: " + str(group))
        #    user.object.groups.add(group)
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        status = thisassessment.status

        owner = form.cleaned_data['owner']
        assessment_name = form.cleaned_data['name']
        if form.cleaned_data['vendor_contact']:
            vendor_POC = form.cleaned_data['vendor_contact']
            #print("form vendor POC:" + str(vendor_POC))
            #print("initial POC:" + str(initial_vendor_POC))
            if str(vendor_POC) != str(initial_vendor_POC):
                subject = assessment_name
                message = "Hello, " + vendor_POC.first_name + \
                  "   You have been identified by " + owner.first_name + " " + owner.last_name + \
                  " as the Vendor Point of Contact for the " + assessment_name
                from_email = owner.email
                to_email = vendor_POC.email
                send_mail(
                    subject,
                    message,
                    from_email,
                    recipient_list = [ to_email ]
                )

        if form.cleaned_data['analyst']:
            assmt_analyst = form.cleaned_data['analyst']

            if str(assmt_analyst) != str(init_analyst):
                subject = assessment_name
                message = "Hello, " + assmt_analyst.first_name + \
                  "   You have been identified by " + owner.first_name + " " + owner.last_name + \
                  " as the Analyst for the " + assessment_name
                from_email = owner.email
                to_email = assmt_analyst.email
                send_mail(
                    subject,
                    message,
                    from_email,
                    recipient_list = [ to_email ]
                )



        return super().form_valid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        assessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        context['initial_analyst'] = assessment.analyst
        context['initial_vendor_contact'] = assessment.vendor_contact
        return context

def send_email_to_vendorPOC(assessment, assessment_status, request):
    if assessment_status == "VENDOR_SUBMIT":
        subject = assessment.name
        message_html = "<html><body>Hello, " + assessment.vendor_contact.first_name
        message_html += "<p>   The questionnaire for " + assessment.name
        message_html += " has been prepared for you by " + assessment.owner.first_name + " " + assessment.owner.last_name + "."
        message_html +=  "<p> Please review the questionnaire and provide your company's answers. "
        message_html += "<p> The questionnaire can be accessed at <a href='" + request.build_absolute_uri('/').strip("/") + "/questionnaire/" + str(assessment.uuid) + "/'>" + str(assessment.name) + " Questionnaire'</a>'"
        from_email = "assessments@venchmark.com"
        to_email = assessment.vendor_contact.email
        send_mail(
            subject,
            message_html,
            from_email,
            recipient_list = [ to_email ]
        )

def send_email_to_analyst(assessment, assessment_status):
    if assessment_status == "IN_ANALYSIS":
        subject = assessment.name
        message = "Hello, " + assessment.analyst.first_name + \
          "   The questionnaire for " + assessment.name + \
          " has been answered and filled out by " + assessment.vendor_contact.first_name + " " + assessment.vendor_contact.last_name + "(" + assessment.vendor_contact.email +")" + ". Their answers are ready for your analysis. \n" + \
          " Please review the questionnaire and provide your analysis of " + assessment.vendor.name + "'s answers. \n" + \
          " The questionnaire can be accessed at <a href=\"analysis/'" + str(assessment.uuid) + "/\" >" + assessment.name + "</a>"
        from_email = "assessments@venchmark.com"
        to_email = str(assessment.analyst.email)
        send_mail(
            subject,
            message,
            from_email,
            recipient_list = [ to_email ]
        )


def FrameworkControlsUploadView(request, uuid):
    #open the file that was uploaded
    #for each line in the file,
    #set the control's fields
    #save the control

    newfile = request.FILES['controls_file']
    print(newfile.path)
    #newdoc.configfilename = ntpath.basename(filepath)

    #UploadControls(thisassessment.uuid)
    return redirect('/controls/' + str(uuid))

def AssessmentStatusChange(request, uuid):
    class AssesmentStatusSwitch:

        def switch(self, status):
            default = "No Status to Change"
            return getattr(self, status, lambda: default)()

        def CREATED(self):
            questionnaire = Questionnaire.objects.get(assessment=thisassessment.uuid)
            PopulateQuestions(thisassessment.uuid, questionnaire.uuid)
            new_status = "QUESTIONNAIRE_REVIEW"
            thisassessment.status = new_status
            thisassessment.save()
            #url = '/questionnaire/' + str(uuid)
            return

        def QUESTIONNAIRE_REVIEW(self):
            new_status = "VENDOR_SUBMIT"
            thisassessment.status = new_status
            thisassessment.save()
            send_email_to_vendorPOC(thisassessment, thisassessment.status, request)
            #url = '/assessments/'
            return

        def VENDOR_SUBMIT(self):
            new_status = "IN_ANALYSIS"
            thisassessment.status = new_status
            thisassessment.save()
            send_email_to_analyst(thisassessment, thisassessment.status)
            #response = redirect('/')
            #url = '/assessments/'
            return

        def IN_ANALYSIS(self):
            new_status = "REPORT_GENERATED"
            thisassessment.status = new_status
            thisassessment.save()
            #send_email_to_owner(thisassessment, thisassessment.status)
            #url = '/report/' + str(uuid)
            #url = '/'
            #return HttpResponseRedirect(url)
            #return response
            return

        def REPORT_GENERATED(self):
            new_status = "ASSESSMENT_COMPLETE"
            thisassessment.status = new_status
            thisassessment.save()
            #send_email_to_owner(thisassessment, thisassessment.status)
            #url = '/report/' + str(uuid)
            #url = '/'
            return

        #def ASSESSMENT_REVIEW(self):
        #    new_status = "ASSESSMENT_COMPLETE"
        #    thisassessment.status = new_status
        #    thisassessment.save()
        #    #send_email_to_owner(thisassessment, thisassessment.status)
            #url = '/report/' + str(uuid)
        #    #url = '/'
        #    return
    #group_required = [ u"owner", u"administrator"]
    #login_url = '/'
    #model = Assessment
    thisassessment = Assessment.objects.filter(uuid=uuid).first()
    #slug_url_kwarg = Assessment.uuid
    #template_name = 'assessments/assessment_update_form.html'
    #form_class = UpdateAssessmentForm
    template_name = "assessments/assessment_list.html"
    url = '/assessments/'
    #def get_queryset(self):
    #    #queryset = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
    #    queryset = Assessment.objects.filter(uuid=uuid).first()
    #    return queryset

    s = AssesmentStatusSwitch()
    s.switch(thisassessment.status)
    print("returning the user to: " + str(url))
    return HttpResponseRedirect(url)

class CloneFramework(View):


    def get(self, request, uuid):
        #uuid_to_clone = self.request.kwargs['uuid']

        def CloneAFramework(uuid):
            #print("cloning ... ")
            framework_to_clone = Framework.objects.filter(uuid=uuid).first()
            print("cloning: " + str(framework_to_clone.name) + " with UUID: " + str(uuid))
            NewFramework = Framework()
            NewFramework.name = framework_to_clone.name + " - clone"
            NewFramework.short_name = framework_to_clone.short_name
            NewFramework.version = framework_to_clone.version
            NewFramework.publication_date = framework_to_clone.publication_date
            NewFramework.source = framework_to_clone.source
            NewFramework.save()
            LatestFramework = Framework.objects.latest('order_id')
            LFUUID = LatestFramework.uuid
            print("UUID to assign to cloned controls ... [for new framework]: " + str(LFUUID))

            CloneFrameworkControls = FrameworkControls.objects.filter(frameworkUUID=uuid)
            print(str(len(CloneFrameworkControls)))

            for control in CloneFrameworkControls:
                print("cloning " + control.subcategoryID)
                NewControl = FrameworkControls()
                NewControl.framework = LatestFramework #ontrol.framework
                NewControl.frameworkUUID = LFUUID
                NewControl.function = control.function
                NewControl.functionID = control.functionID
                NewControl.category = control.category
                NewControl.categoryID = control.categoryID
                NewControl.category_statement = control.category_statement
                NewControl.subcategory = control.subcategory
                NewControl.subcategoryID = control.subcategoryID
                #NewControl.control_statement = control.control_statement
                NewControl.default_question = control.default_question
                NewControl.reference = control.reference
                NewControl.save()


        CloneAFramework(uuid)
        return HttpResponseRedirect(reverse_lazy('framework_list'))


def VendorQuestionListApproved(request, uuid):
    login_url = '/'
    thisquestionnaire = Questionnaire.objects.filter(assessment=uuid).first()
    thisassessment = Assessment.objects.filter(uuid=uuid).first()
    #AssessmentStatusChange(request,uuid)
    response = redirect('/assessment/' + str(uuid) + "/status-change/")
    return response



class AnalysisListView2(GroupRequiredMixin, ListView):
    group_required = [u"administrator", u"owner", u"analyst"]
    login_url = '/'
    model = Assessment
    template_name = "assessments/analysis_list.html"
    #queryset = Questionnaire.objects.all()
    paginate_by = 10
    #queryset = Assessment.objects.filter(analyst=analyst)
    #fields = ['uuid', 'Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    fields = ['uuid', 'Name' ]

    #def get_object(self, queryset=None):
        #obj = User.objects.filter(uuid=self.kwargs['uuid']).first()
        #myAssessment = Assessment.objects.get(pk=self.kwargs['uuid'],analyst=self.kwargs['analyst'])
    #    myAssessment = Assessment.objects.get(analyst=self.kwargs['analyst'])
        #print(myUser.first_name + " " + myUser.last_name)
        #groups = myUser.groups.all()
        #for group in groups:
        #    print(group)
        #print(groups)
        #print("user belongs " + str(groups.count()) + " groups.")
        #for group in groups:
    #        print("retrieved group" + str(group))
    #    return myAssessment

    def get_queryset(self):
        #return Assessment.objects.filter(analyst=self.kwargs['analyst'])
        #print(self.request.user)
        return Assessment.objects.filter(analyst=self.request.user)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        #user = self.request.user
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

    def get_initial(self):
    #return {'Owner': self.kwargs['request.user']}
        return {'analyst': self.request.user}

class AnalysisDetailView(GroupRequiredMixin, FormView):
    group_required = [ u"administrator", u"owner", u"analyst" ]
    #authenticated_redirect_url = u"/"
    #login_url = '/users/login/'
    login_url = '/'
    redirect_unauthenticated_users = True
    slug_url_kwarg = Assessment.uuid
    model = Question
    #if current_user.groups.filter(name = u"vendor").exists():
    form_class = AnalysisQuestionnaireForm
    #elif current_user.groups.filter(name = u"analyst").exists():
    #    form_class = AnalystContactQuestionnaireForm
    #elif django_user.groups.filter(name = u"adminstrator").exists():
    #    form_class = AdministratorContactQuestionnaireForm
    template_name = "assessments/analysis_list.html"
    #context_object_name = 'Framework_list'   # your own name for the list as a template variable
    #queryset = Question.objects.filter(questionnaire_id=self.kwargs['uuid'])
    fields = "__all__"
    #fields = ['Name', 'ShortName', 'Version', 'PublishDate', 'Source']
    #template_name = 'Frameworks/Frameworks_list.html'  # Specify your own template name/location
    paginate_by = 10

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        questions = Question.objects.filter(assessment=self.kwargs['uuid']).all()
        context['questions'] = questions
        context['uuid'] = self.kwargs['uuid']
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class ReportListView(GroupRequiredMixin, ListView):
    group_required = [u"administrator", u"owner", u"analyst"]
    login_url = '/'
    model = Report
    template_name = "assessments/reports_list.html"
    #queryset = Questionnaire.objects.all()

    #queryset = Assessment.objects.filter(analyst=analyst)
    #fields = ['uuid', 'Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    fields = "__all__" #['uuid', 'Name' ]

    def get_object(self, queryset=None):
        #obj = User.objects.filter(uuid=self.kwargs['uuid']).first()
        #myReports = Report.objects.get(pk=self.kwargs['uuid'],analyst=self.kwargs['analyst'])
        #Q(vendor_contact=self.request.user) | Q(analyst=self.request.user) | Q(owner=self.request.user)
        myReports = Report.objects.filter(Q(analyst=self.request.user) | Q(owner=self.request.user))
        #print(myUser.first_name + " " + myUser.last_name)
        #groups = myUser.groups.all()
        #for group in groups:
        #    print(group)
        #print(groups)
        #print("user belongs " + str(groups.count()) + " groups.")
        #for group in groups:
    #        print("retrieved group" + str(group))
        return myAssessment

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        #user = self.request.user
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

    def get_initial(self):
    #return {'Owner': self.kwargs['request.user']}
        return {'analyst': self.request.user, 'owner': self.request.user}

def GenerateReport(request, uuid):
        questionnaire = Questionnaire.objects.filter(assessment=uuid).first()
        assessment = Assessment.objects.filter(uuid=uuid).first()
        report_check = Report.objects.filter(assessment_id=uuid).first()
        user = request.user

        if not report_check:
            NewReport = Report()
            NewReport.name = assessment.name + " Report"
            NewReport.publication_date = datetime.today()
            NewReport.questionnaire = questionnaire
            NewReport.assessment = assessment
            NewReport.executive_summary = ""
            NewReport.executive_compliance_summary = ""
            NewReport.executive_risk_assessment_summary = ""
            NewReport.nc_controls = 0
            NewReport.c_controls = 0
            NewReport.author = user
            NewReport.save()
        generatedReport = Report.objects.filter(assessment_id=uuid).first()
        url = request.META.get('HTTP_REFERER')
        url = '/report/' + str(generatedReport.uuid)
        return HttpResponseRedirect(url)



class ReportGenerateView(GroupRequiredMixin, FormView):
    group_required = [ u"owner", u"administrator", u"analyst" ]
    login_url = '/'
    #success_url = '/assessment/'+str(self.kwargs['uuid'])+'/'
    #template_name = 'Framework_new.html'
    template_name = 'assessments/NewReportForm.html'
    #author = self.kwargs['request_user']
    form_class = CreateReportForm

    def form_valid(self, form):
        # This method is called when valid form data has been POSTed.
        # It should return an HttpResponse.
        form.GenerateReport(self.kwargs['uuid'], self.request.user)
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        return super().form_valid(form)

    def get_success_url(self):
            return '/analysis/'+str(self.kwargs['uuid'])+'/'   #reverse('assessment_detail'+'/'+str(self.kwargs['uuid'])+'/')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #print("frameworkUUID = " + str(self.kwargs['uuid']))

        rassessment = Assessment.objects.filter(uuid=self.kwargs['uuid']).first()
        rquestionnaire = Questionnaire.objects.filter(assessment=rassessment)
        #frameworks = rassessment.frameworks.all()
        context['questionnaire'] = rquestionnaire
        context['assessment'] = rassessment
        #context['frameworks'] = frameworks
        #context['controls'] = controls
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

    #def get_initial(self):
        #return {'Owner': self.kwargs['request.user']}
    #    return {'author': self.request.user}

class ReportEditView(GroupRequiredMixin, DetailView):
    group_required = [ u"analyst", u"owner", u"administrator"]
    login_url = '/'
    template_name = "assessments/report_edit.html"
    model = Report
    slug_url_kwarg = Report.uuid
    fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    def get_object(self, queryset=None):
        obj = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Report.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        #print("Questions count: " + str(questions.count()))
        context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class ReportDetailView(GroupRequiredMixin, DetailView):
    group_required = [ u"analyst", u"owner", u"administrator"]
    login_url = '/'
    template_name = "assessments/report_view2.html"
    model = Report
    slug_url_kwarg = Report.uuid
    fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    def get_object(self, queryset=None):
        obj = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Report.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        nc_questions = Question.objects.filter(assessment=report.assessment.uuid, compliance=False)
        #print("Questions count: " + str(questions.count()))
        context['questions'] = questions
        context['nc_questions'] = nc_questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context



class ReportDetailPdf(View):

    #slug_url_kwarg = Report.uuid

    #def get(request, uuid):
    def get(self, request, uuid):
        #uuid = self.kwargs['uuid']
        #sales = Sales.objects.all()
        #today = timezone.now()
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        #print("Questions count: " + str(questions.count()))
        #context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        #context['token'] = token
        params = {
            'report': report,
            'questions': questions,
            'token': token
        }
        #return Render.render('assessments/report_pdf_detail.html', params)
        return Render.render('assessments/report_view2.html', params)

class ReportDetailPdf2(View):
    def get(self, request, uuid):
        response = HttpResponse(content_type='application/pdf')
        d = datetime.today().strftime('%Y-%m-%d')
        # Create a file-like buffer to receive PDF data.
        buffer = BytesIO()

        # Create the PDF object, using the buffer as its "file."
        p = canvas.Canvas(buffer, pagesize=letter)

        # data to print
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        questions = Question.objects.filter(assessment=report.assessment.uuid)

        token = Token.objects.get(user=self.request.user)

        # Draw things on the PDF. Here's where the PDF generation happens.
        # See the ReportLab documentation for the full list of functionality.
        p.setFont("Helvetica",15,leading=None)
        x1 = 0
        y1 = 0

        p.drawString(100, 100, report.executive_summary)
        p.drawString(x1, y1-20, report.executive_findings)
        p.drawString(x1, y1-20, report.executive_compliance_summary)
        p.drawString(x1, y1-20, report.executive_risk_assessment_summary)

        # Close the PDF object cleanly, and we're done.
        p.showPage()
        p.save()

        # FileResponse sets the Content-Disposition header so that browsers
        # present the option to save the file.
        buffer.seek(0)
        return FileResponse(buffer, as_attachment=True, filename='{d}-report.pdf')

def render_to_pdf(template_src, context_dict):
    template = get_template(template_src)
    #context = Context(context_dict)#
    context = dict(context_dict)
    html  = template.render(context)
    print(html)
    #result = StringIO.StringIO()
    result = StringIO()

    #pdf = pisa.pisaDocument(StringIO.StringIO(html.encode("ISO-8859-1")), result)
    #pdf = pisa.pisaDocument(StringIO(html.encode("ISO-8859-1")), result)
    #pdf = pisa.pisaDocument(StringIO(html.encode("UTF-8")), result)
    pdf = pisa.pisaDocument(BytesIO(html.encode("UTF-8")), result)
    if not pdf.err:
        return HttpResponse(result.getvalue(), content_type='application/pdf')
    return HttpResponse('We had some errors<pre>%s</pre>' % escape(html))

class ReportDetailPdf3(View):
    def get(self, request, uuid):
        #Retrieve data or whatever you need
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #questions = Question.objects.filter(assessment=report.assessment.uuid)
        token = Token.objects.get(user=self.request.user)
        return render_to_pdf(
                'assessments/report_view2.html',
                {
                    'pagesize':'A4',
                    'report': report,
                    #'questions': questions,
                    'token': token,
                }
            )


class ReportPDF(PDFTemplateView):
    #print("ReportPDF")
    d = datetime.today().strftime('%Y-%m-%d')
    filename = 'report.pdf'
    template_name = 'assessments/report_view3.html'
    cmd_options = {
        'margin-top': 3,
    }

class MyPdfView(View):
    def get(self, request, *args, **kwargs):
        report_uuid = self.kwargs['uuid']
        file_name = BASE_DIR + '/assessments/testreport.pdf'
        filefrom = BASE_DIR + '/assessments/convertpdf.js'
        url = 'http://' + request.get_host() + '/report/pdf/' + str(report_uuid) + '/'
        print(url)
        #url = 'http'
        external_process = Popen(["/usr/local/bin/phantomjs", filefrom, url, file_name], stdout=PIPE, stderr=STDOUT)
        external_process.wait()
        return_file = FileWrapper(open(file_name, 'r'))
        download_file_name = 'report'
        response = StreamingHttpResponse(return_file, content_type='application/force-download')
        response['Content-Disposition'] = 'attachment; filename= %s.pdf' % download_file_name
        return response

def link_callback(uri, rel):
    """
    Convert HTML URIs to absolute system paths so xhtml2pdf can access those
    resources
    """
    result = finders.find(uri)
    if result:
            if not isinstance(result, (list, tuple)):
                    result = [result]
            result = list(os.path.realpath(path) for path in result)
            path=result[0]
    else:
            sUrl = settings.STATIC_URL        # Typically /static/
            sRoot = settings.STATIC_ROOT      # Typically /home/userX/project_static/
            mUrl = settings.MEDIA_URL         # Typically /media/
            mRoot = settings.MEDIA_ROOT       # Typically /home/userX/project_static/media/

            if uri.startswith(mUrl):
                    path = os.path.join(mRoot, uri.replace(mUrl, ""))
            elif uri.startswith(sUrl):
                    path = os.path.join(sRoot, uri.replace(sUrl, ""))
            else:
                    return uri

    # make sure that file exists
    if not os.path.isfile(path):
            raise Exception(
                    'media URI must start with %s or %s' % (sUrl, mUrl)
            )
    return path

#---------- such a cluster
class ReportPDF4(View):
    def get(self, request, *args, **kwargs):
        template_path = 'assessments/report_view2.html'
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        token = Token.objects.get(user=self.request.user)
        context = {'report': report, 'token': token}
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        # find the template and render it.
        template = get_template(template_path)
        html = template.render(context)

        # create a pdf
        pisa_status = pisa.CreatePDF(
           html, dest=response, link_callback=link_callback)
        # if error then show some funy view
        if pisa_status.err:
           return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response


class ReportDetailPDFView(GroupRequiredMixin, DetailView):
    group_required = [ u"analyst", u"owner", u"administrator"]
    login_url = '/'
    template_name = "assessments/report_view_pdf.html"
    model = Report
    slug_url_kwarg = Report.uuid
    #fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']

    def get(self, request, *args, **kwargs):
        template_path = 'assessments/report_view_pdf.html'
        report_uuid = self.kwargs['uuid']
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #print(report.executive_summary)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        nc_questions = Question.objects.filter(assessment=report.assessment.uuid, compliance=False)
        token = Token.objects.get(user=self.request.user)
        #context = {'uuid': report_uuid, 'report': report, 'token': token}
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        # find the template and render it.
        template = get_template(template_path)
        #print(self.get_context_data())
        #context = self.get_context_data()
        pdf_html = template.render({'uuid': report_uuid, 'report': report, 'token': token, 'questions': questions, 'nc_questions': nc_questions  })
        #clean_html = format_html(pdf_html)
        print(pdf_html)

        # create a pdf
        pisa_status = pisa.CreatePDF(pdf_html, dest=response, link_callback=link_callback)
        # if error then show some funy view
        if pisa_status.err:
           return HttpResponse('We had some errors <pre>' + html + '</pre>')
        return response

    def get_object(self, queryset=None):
        obj = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Report.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        #print("Questions count: " + str(questions.count()))
        context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context

class ReportDetailDOCXView(GroupRequiredMixin, DetailView):
    group_required = [ u"analyst", u"owner", u"administrator"]
    login_url = '/'
    template_name = "assessments/report_view_docx.html"
    model = Report
    slug_url_kwarg = Report.uuid
    #fields = ['Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']

    def get(self, request, *args, **kwargs):
        template_path = 'assessments/report_view_docx.html'
        report_uuid = self.kwargs['uuid']
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #print(report.executive_summary)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        token = Token.objects.get(user=self.request.user)
        #context = {'uuid': report_uuid, 'report': report, 'token': token}
        # Create a Django response object, and specify content_type as pdf
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="report.pdf"'
        # find the template and render it.
        template = get_template(template_path)
        #print(self.get_context_data())
        #context = self.get_context_data()
        html = template.render({'uuid': report_uuid, 'report': report, 'token': token})
        title = report.name + str(" Report")
        #print(html)

        # create a DOCX
        try:
            buf = html2docx(html, title=title)
            doc = docx.Document(buf)

            with open("my.docx", "wb") as fp:
                fp.write(buf.getvalue())
        except Exception as e:
            print(e)
            return HttpResponse('We had some errors <pre>' + html + '</pre>')

        return response

def replaceit(string):
    cleanstring2 = string.replace("&lt", "<")
    cleanstring = cleanstring2.replace("&gt", ">")
    return(cleanstring)

def ReportDetailDOCXView2(request, uuid):
    #form = ReportGenForm(request.POST or None)

    #if form.is_valid():
        #doctype = form.cleaned_data['format']str(settings.MEDIA_ROOT) +
    report = Report.objects.filter(uuid=uuid).first()
    #report_dict = {}
    #compliance_image = "/media/uploads/" + str(uuid) + "/" + "ComplianceChart.png"
    #print(compliance_image)
    clean_report_exec_sum = html.unescape(report.executive_summary)
    report_dict = { 'name': report.name, 'executive_summary': clean_report_exec_sum, 'executive_findings': report.executive_findings, 'executive_compliance_summary': report.executive_compliance_summary, 'executive_risk_assessment_summary': report.executive_risk_assessment_summary, 'executive_compliance_chart': report.executive_compliance_chart, 'executive_risk_chart': report.executive_risk_chart }
    #report_dict = { 'name': report.name, 'executive_summary': report.executive_summary, 'executive_findings': report.executive_findings, 'executive_compliance_summary': report.executive_compliance_summary, 'executive_risk_assessment_summary': report.executive_risk_assessment_summary, 'executive_compliance_chart': compliance_image, 'executive_risk_chart': report.executive_risk_chart }

    doctype = 'docx'
    filename = fill_template(
        #'assessments/report.odt', form.cleaned_data,
        'assessments/report.odt', report_dict,
        output_format=doctype)
    visible_filename = 'report.{}'.format(doctype)

    return FileResponse(filename, visible_filename)
#else:
#    return HttpResponseRedirect('/reports/')

#    return response



    def get_object(self, queryset=None):
        obj = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Report.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        #print("Questions count: " + str(questions.count()))
        context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context



def ReportDetailDOCXView3(request, uuid):

    def report_based_upload_to(uuid, filename):
        return "uploads/{}/{}".format(uuid, filename)

    def change_orientation():
        current_section = doc.sections[-1]
        new_width, new_height = current_section.page_height, current_section.page_width
        new_section = doc.add_section(WD_SECTION.NEW_PAGE)
        new_section.orientation = WD_ORIENT.LANDSCAPE
        new_section.page_width = new_width
        new_section.page_height = new_height

        return new_section

    def set_repeat_table_header(row):
        """ set repeat table row on every new page
        """
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        tblHeader.set(qn('w:val'), "true")
        trPr.append(tblHeader)
        return row

    def create_presigned_url(bucket_name, object_name, expiration=3600):
        """Generate a presigned URL to share an S3 object

        :param bucket_name: string
        :param object_name: string
        :param expiration: Time in seconds for the presigned URL to remain valid
        :return: Presigned URL as string. If error, returns None.
        """

        # Generate a presigned URL for the S3 object
        s3_client = boto3.client('s3')
        try:
            response = s3_client.generate_presigned_url('get_object',
                                                        Params={'Bucket': bucket_name,
                                                                'Key': object_name},
                                                        ExpiresIn=expiration)
        except ClientError as e:
            logging.error(e)
            return None

        # The response contains the presigned URL
        return response

    report = Report.objects.filter(uuid=uuid).first()
    questions = Question.objects.filter(assessment=report.assessment.uuid)
    nc_questions = Question.objects.filter(assessment=report.assessment.uuid, compliance=False)
    #report_dict = {}
    #compliance_image = "/media/uploads/" + str(uuid) + "/" + "ComplianceChart.png"
    #print(compliance_image)
    #clean_report_exec_sum = html.unescape(report.executive_summary)
    #report_dict = { 'name': report.name, 'executive_summary': clean_report_exec_sum, 'executive_findings': report.executive_findings, 'executive_compliance_summary': report.executive_compliance_summary, 'executive_risk_assessment_summary': report.executive_risk_assessment_summary, 'executive_compliance_chart': report.executive_compliance_chart, 'executive_risk_chart': report.executive_risk_chart }
    # Create an instance of a word document
    #doc.save(filename)
    clean_html_exec_sum = html.unescape(report.executive_summary)
    clean_html_exec_findings = html.unescape(report.executive_findings)
    clean_html_exec_compliance_sum = html.unescape(report.executive_compliance_summary)
    clean_html_exec_risk_assessment = html.unescape(report.executive_risk_assessment_summary)

    doc = docx.Document()


    # Add a Title to the document
    doc.add_heading(report.name, 0)

    # Adding paragraph with new font Style
    doc.add_heading('Executive Summary', 3)
    html2docx(clean_html_exec_sum, "report", doc)
    # Add a field from the DB
    #print(html.escape(report.executive_summary))
    #paragraph = add_text(html.unescape(report.summary))
    #clean_html_exec_sum = html.unescape(report.executive_summary)
    #parser = html2docx(clean_html_exec_sum, "Executive Summary")
    #thing_to_print = HTML2Docx.add_text(clean_html_exec_sum)
    #print(str(thing_to_print))
    #print(str(parser))
    #print(parser)
    #parser.feed(content.strip())
    #print(parser)
    #doc.add_paragraph(
    #    str(parser))

    # Add some spacing
    doc.add_paragraph()
    doc.add_paragraph()


    # Adding paragraph with new font Style
    doc.add_heading('Findings', 3)

    # Add a field from the DB
    html2docx(clean_html_exec_findings, "report", doc)


    # Add some spacing
    doc.add_paragraph()
    doc.add_paragraph()

    # Adding paragraph with new font Style
    doc.add_heading('Security Controls Compliance', 3)

    # Add a field from the DB
    html2docx(clean_html_exec_compliance_sum, "report", doc)


    # Add some spacing
    doc.add_paragraph()
    doc.add_paragraph()


    doc.add_picture(report.executive_compliance_chart, width=Inches(4.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding a page break
    doc.add_page_break()

    # Adding paragraph with new font Style
    doc.add_heading('Risk Assessment', 3)

    # Add a field from the DB
    html2docx(clean_html_exec_risk_assessment, "report", doc)


    doc.add_picture(report.executive_risk_chart, width=Inches(4.0))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Adding a page break
    doc.add_page_break()

    # Add a Title to the document
    doc.add_heading('Security Controls Compliance', 0)

    # Creating a table object
    table = doc.add_table(rows=0, cols=5, style='Table Grid')
    row = table.add_row()
    set_repeat_table_header(row)

    # Adding heading in the 1st row of the table
    row = table.rows[0].cells

    row[0].text = 'Control'
    row[1].text = 'Question'
    row[2].text = 'Answer'
    row[3].text = 'Compliance'
    row[4].text = 'Notes'
    #set_repeat_table_header(row)

    # Adding data from the list to the table
    #for id, name in data:
    for question in questions:

        # Adding a row and then adding data in it.
        row = table.add_row().cells
        # Converting id to string as table can only take string input
        row[0].text = str(question.control.subcategoryID)
        row[1].text = str(question.question)
        row[2].text = str(question.answer)
        row[3].text = str(question.compliance)
        row[4].text = str(question.notes)

    # Adding a page break
    doc.add_page_break()


    change_orientation()

    if len(nc_questions) >= 1:
        doc.add_heading('Risk Assessment - Non-Compliant Controls', 3)
        # Adding data from the list to the table
        # Creating a table object
        table = doc.add_table(rows=0, cols=8, style='Table Grid')
        row = table.add_row()
        set_repeat_table_header(row)
        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = 'Control'
        row[1].text = 'Question'
        row[2].text = 'Answer'
        row[3].text = 'Compliance'
        row[4].text = 'Notes'
        row[5].text = 'Likelihood of Exploit'
        row[6].text = 'Impact'
        row[7].text = 'Risk Rating'

        set_repeat_table_header(row)


        for nc_question in nc_questions:

            # Adding a row and then adding data in it.
            row = table.add_row().cells
            # Converting id to string as table can only take string input
            row[0].text = str(nc_question.control.subcategoryID)
            row[1].text = str(nc_question.question)
            row[2].text = str(nc_question.answer)
            row[3].text = str(nc_question.compliance)
            row[4].text = str(nc_question.notes)
            row[5].text = str(nc_question.likelihood)
            row[6].text = str(nc_question.impact)
            row[7].text = str(nc_question.risk_rating)
    else:
        doc.add_heading('Risk Assessment - Non-Compliant Controls', 3)
        # Adding data from the list to the table
        # Creating a table object
        table = doc.add_table(rows=0, cols=8, style='Table Grid')
        row = table.add_row()
        set_repeat_table_header(row)

        # Adding heading in the 1st row of the table
        row = table.rows[0].cells
        row[0].text = 'Control'
        row[1].text = 'Question'
        row[2].text = 'Answer'
        row[3].text = 'Compliance'
        row[4].text = 'Notes'
        row[5].text = 'Likelihood of Exploit'
        row[6].text = 'Impact'
        row[7].text = 'Risk Rating'

        # Add some spacing
        doc.add_paragraph()
        doc.add_paragraph("No Controls were shown to bring risk to the company.")


    # Now save the document to a location
    filename = report.assessment.name + " Report.docx"



    temp_filename = '/app/tmp/' + filename
    doc.save(temp_filename)


    AWS_S3_UPLOAD_PATH = report_based_upload_to(report.uuid, filename)
    #file_URL = AWS_S3_BUCKET_URL + AWS_S3_UPLOAD_PATH
    import boto3
    s3 = boto3.client('s3')
    s3.upload_file(temp_filename, settings.AWS_STORAGE_BUCKET_NAME, AWS_S3_UPLOAD_PATH)

    #bucket = s3.Bucket(settings.AWS_S3_BUCKET_NAME)
    #bucket = s3.Bucket(settings.AWS_STORAGE_BUCKET_NAME)


    #source_stream = StringIO(doc)
    #bucket_path = report_based_upload_to(report.uuid, filename)
    print("AWS S3 Upload Path: " + AWS_S3_UPLOAD_PATH)
    #s3.upload_file(temp_filename, settings.AWS_STORAGE_BUCKET_NAME, AWS_S3_UPLOAD_PATH)
    #body=obj.get()['Body'].read()
    #buffer = BytesIO()
    #buffer.write(doc)
    #print(buffer)
    #text = docx2txt.process(buffer)

    #bucket.put_object(Key=bucket_path, Body=buffer)
    #bucket.put_object(Key=AWS_S3_UPLOAD_PATH, Body=buffer)


    #doc.save(file_URL)
    #visible_filename = 'report.{}'.format(doctype)
    visible_filename = filename
    #report_file = open()
    #url = create_presigned_url(settings.AWS_STORAGE_BUCKET_NAME, filename)
    url = create_presigned_url(settings.AWS_STORAGE_BUCKET_NAME, AWS_S3_UPLOAD_PATH)
    print(url)
    if url is not None:
        response = requests.get(url)
    return HttpResponseRedirect(url)
    #return FileResponse(url, visible_filename)
    #return FileResponse(response)
    #return HttpResponse(response)
    #return StreamingHttpResponse(response)

#else:
#    return HttpResponseRedirect('/reports/')

#    return response



    def get_object(self, queryset=None):
        obj = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        return obj

    def get_queryset(self):
        print("self kwarg id:" + self.kwargs['uuid'])
        queryset = Report.objects.filter(uuid=self.kwargs['uuid']).order_by('id')
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        report = Report.objects.filter(uuid=self.kwargs['uuid']).first()
        #assessment = Assessment.objects.filter(uuid=report.assessment.uuid)
        questions = Question.objects.filter(assessment=report.assessment.uuid)
        #print("Questions count: " + str(questions.count()))
        context['questions'] = questions
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context


def check_membership(request, groups):
    """ Check required group(s) """
    if request.user.is_superuser:
        return True
    user_groups = request.user.groups.values_list("name", flat=True)
    print(user_groups)
    return set(groups).intersection(set(user_groups))


@csrf_exempt
@login_required
def answer_update(request, uuid):
    #if not request.user.has_perm('backend.add_artist'):
    #if not check_membership(request, r"vendor"): #request.user.has_role('vendor'):
    #    raise PermissionDenied

    # Either render only the modal content, or a full standalone page
    #if request.is_ajax():
    #    template_name = 'frontend/includes/generic_form_inner.html'
    #else:
    #    template_name = 'frontend/includes/generic_form.html'

    object = None
    if request.method == 'POST':
        form = AnswerUpdateForm(data=request.POST)
        if form.is_valid():

            #questionnaire = Questionnaire.objects.get(assessment=thisassessment.uuid)
            #PopulateQuestions(thisassessment.uuid, questionnaire.uuid)
            #new_status = "QUESTIONNAIRE_REVIEW"
            #thisassessment.status = new_status
            #thisassessment.save()

            question = Question.objects.filter(uuid=uuid).first()
            question.answer = form.cleaned_data['answer']
            question.save()
            #question.answer = request.answer  #self.kwargs['uuid']).first()
            #object = form.save()
            if not request.is_ajax():
                # reload the page
                #next = request.META['PATH_INFO']
                return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
                #return HttpResponseRedirect(next)
            # if is_ajax(), we just return the validated form, so the modal will close
    else:
        form = AnswerUpdateForm()

    #next = request.META['PATH_INFO']
    #return HttpResponseRedirect(next)
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
    #return render(request, template_name, {
    #    'form': form,
    #    'object': object,
    #})
    #framework_to_clone=Framework.objects.filter(uuid=uuid)
    #NewFramework = Framework()
    #CloneFramework(uuid)
    #return HttpResponseRedirect(reverse_lazy('framework_list'))):
    #question_uuid = request.GET['uuid']
    #obj = models.Post.objects.get(pk=pk)
    #question = Question.objects.filter(uuid=self.kwargs['uuid'])  #.first()
    #obj.upvotes = obj.upvotes +1

    #question.save()
    #response_data = {'success':1}
    #return HttpResponse(json.dumps(response_data), content_type="application/json")

@csrf_exempt
@login_required
def question_edit_update(request, uuid):
    #if not request.user.has_perm('backend.add_artist'):
    #if not check_membership(request, r"vendor"): #request.user.has_role('vendor'):
    #    raise PermissionDenied

    # Either render only the modal content, or a full standalone page
    #if request.is_ajax():
    #    template_name = 'frontend/includes/generic_form_inner.html'
    #else:
    #    template_name = 'frontend/includes/generic_form.html'

    object = None
    if request.method == 'POST':
        form = QuestionEditUpdateForm(data=request.POST)
        if form.is_valid():

            question = Question.objects.filter(uuid=uuid).first()
            question.question = form.cleaned_data['question']
            question.save()
            #question.answer = request.answer  #self.kwargs['uuid']).first()
            #object = form.save()
            if not request.is_ajax():
                # reload the page
                #next = request.META['PATH_INFO']
                return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
                #return HttpResponseRedirect(next)
            # if is_ajax(), we just return the validated form, so the modal will close
    else:
        form = QuestionEditUpdateForm()

    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))


@csrf_exempt
@login_required
def report_update(request, uuid):

    #object = None
    if request.method == 'POST':
        form = ReportUpdateForm(data=request.POST)
        #print(form)
        #print(form)
        #print("passed: " + str(form.cleaned_data['id_executive_summary']))
        #'executive_summary',
        #'executive_findings',
        #'executive_compliance_summary',
        #'executive_risk_assessment_summary',
        #'nc_controls',
        #'c_controls',
        #FieldsFormSet = inlineformset_factory(executive_summary, executive_findings, executive_compliance_summary, executive_risk_assessment_summary)
        #formset=FielsdFormSet(request.POST)
        #print(form.errors)
        #print(form.cleaned_data['executive_summary'])
        #print(form.cleaned_data['executive_findings'])
        #print(form.cleaned_data['executive_compliance_summary'])
        #print(form.cleaned_data['executive_risk_assessment_summary'])
        #print(request.POST['executive_compliance_chart'])
        #print(request.POST['executive_risk_chart'])
        #print("compliance URL: " + str(form.cleaned_data['executive_compliance_chart']))
        #print("risk URL: " + str(form.cleaned_data['executive_risk_chart']))
        if form.is_valid():

            #questionnaire = Questionnaire.objects.get(assessment=thisassessment.uuid)
            #PopulateQuestions(thisassessment.uuid, questionnaire.uuid)
            #new_status = "QUESTIONNAIRE_REVIEW"
            #thisassessment.status = new_status
            #thisassessment.save()

            thisreport = Report.objects.filter(uuid=uuid).first()


            ## save compliance chart first_name
            ## read in the compliance chart toURL data from request
            #image_data = form.cleaned_data['executive_compliance_chart']
            image_data = request.POST['executive_compliance_chart']
            #print("SURVEY SAYS!: " + str(image_data))
            #get extension
            format, imgstr = image_data.split(';base64,')
            #  print("format", format)
            ext = format.split('/')[-1]

            ## decode base64 encoded file
            data = ContentFile(base64.b64decode(imgstr))

            #set filename
            file_name = "ComplianceChart." + ext

            filepath = os.path.join(str(settings.MEDIA_ROOT), 'uploads', str(uuid), file_name)
            #filepath +=
            # Handle errors while calling os.remove()
            try:
                os.remove(filepath)
            except:
                print("Error while deleting file ", filepath)
            #save this to field - this needs modification
            thisreport.executive_compliance_chart.save(file_name, data, save=True) # image is User's model field


            ## save compliance chart first_name
            ## read in the compliance chart toURL data from request
            #image_data = form.cleaned_data['executive_compliance_chart']
            image_data = request.POST['executive_risk_chart']
            #print("SURVEY SAYS!: " + str(image_data))
            #get extension
            format, imgstr = image_data.split(';base64,')
            #  print("format", format)
            ext = format.split('/')[-1]

            ## decode base64 encoded file
            data = ContentFile(base64.b64decode(imgstr))

            #set filename
            file_name = "RiskChart." + ext
            filepath = os.path.join(str(settings.MEDIA_ROOT), 'uploads', str(uuid), file_name)
            # Handle errors while calling os.remove()
            try:
                os.remove(filepath)
            except:
                print("Error while deleting file ", filepath)
            #save this to field - this needs modification
            thisreport.executive_risk_chart.save(file_name, data, save=True) # image is User's model field


            # save rest of report
            thisreport.executive_summary = form.cleaned_data['executive_summary']
            #print("form received exec sum: " + str(thisreport.executive_summary))
            thisreport.executive_findings = form.cleaned_data['executive_findings']
            #print("form received exec findings: " + str(thisreport.executive_findings))
            #thisreport.executive_findings = form.cleaned_data['id_executive_findings']
            thisreport.executive_compliance_summary = form.cleaned_data['executive_compliance_summary']
            #print("form received exec compliance summary: " + str(thisreport.executive_compliance_summary))
            thisreport.executive_risk_assessment_summary = form.cleaned_data['executive_risk_assessment_summary']
            #print("form received exec risk assessment: " + str(thisreport.executive_risk_assessment_summary))
            #thisreport.executive_compliance_chart = form.cleaned_data['executive_compliance_chart']
            #thisreport.executive_risk_chart = form.cleaned_data['executive_risk_chart']
            thisreport.save()

            #report = Report.objects.filter(uuid=uuid).first()
            #report.executive_summary = form.cleaned_data['executive_summary']
            #print("stored report exec sum: (" + str(report.executive_summary) + ")")
            #'executive_summary',
            #'executive_findings',
            #'executive_compliance_summary',
            #'executive_risk_assessment_summary',
            #'nc_controls',
            #'c_controls',
            #question.answer = request.answer  #self.kwargs['uuid']).first()
            #object = form.save()
            if not request.is_ajax():
                # reload the page
                #next = request.META['PATH_INFO']
                #return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
                return HttpResponseRedirect('/reports/')
                #return HttpResponseRedirect(next)
            # if is_ajax(), we just return the validated form, so the modal will close
    else:
        #print(form)
        #errors = formset.errors
        print(form.errors)
        form = ReportUpdateForm()


    #return HttpResponseRedirect(request.META.get('HTTP_REFERER'))
    return HttpResponseRedirect('/reports/')


def UploadDocument(request, uuid):
    # Handle file upload
    if request.method == 'POST':
        #form = DocumentForm(request.POST, request.FILES)
        form = AnswerDocumentUpdateForm(request.POST, request.FILES)
        print(request.POST)
        print(request.FILES)
        if form.is_valid():
            newdoc = Document(document = request.FILES['document'])
            which_assessment = Assessment.objects.filter(uuid=request.POST['assessment']).first()
            newdoc.assessment = which_assessment
            #question = Question.objects.filter(uuid=request.POST['evidence_q_uuid']).first()
            #newdoc.question = question
            newdoc.name = request.POST['document_name']
            newdoc.description = request.POST['document_description']
            newdoc.save()

            question = Question.objects.filter(uuid=request.POST['question']).first()
            question.document = newdoc #form.cleaned_data['answer']
            #question.document_name = form.cleaned_data['document_name']
            #question.document_description = form.cleaned_data['document_description']
            question.save()
            #newdoc.name = request.cleaned_data['name']

            #newdoc.save()
            #newdoc.configfileid = "{0:0>3}".format(newdoc.id)
            #filepath = newdoc.configfile.path
            #newdoc.configfilename = ntpath.basename(filepath)
            #print(filepath)
            #fileurl = os.path.join(settings.BASE_DIR, 'userfiles/' + newdoc.docfile.name)
            #fileurl = "user" + fileurl
            #parse = CiscoConfParse(newdoc.docfile.url)
            #newdoc.configfilepath = filepath
            #parse = CiscoConfParse(filepath)
    return HttpResponseRedirect(request.META.get('HTTP_REFERER'))

class DocumentCreateView(BSModalCreateView):
    template_name = 'assessments/create_document.html'
    form_class = BSModAnswerDocumentUpdateForm
    success_message = 'Success: Document was added.'
    #success_url = reverse_lazy('index')
    def form_valid(self, form):
        print("Form is valid")
        self.success_url = self.request.META.get('HTTP_REFERER') #self.request.POST.get('previous_page')
        return super().form_valid(form)
    #success_url = HttpResponseRedirect(self.request.META.get('HTTP_REFERER'))

#views go here
class QuestionnaireListView(GroupRequiredMixin, ListView):
    group_required = [u"administrator", u"owner", u"analyst", u"vendor"]
    login_url = '/'
    model = Assessment
    template_name = "assessments/questionnaire_list.html"
    #queryset = Questionnaire.objects.all()

    #queryset = Assessment.objects.filter(analyst=analyst)
    #fields = ['uuid', 'Name', 'Vendor', 'Owner', 'StartDate', 'CompleteDate', 'Status', 'Framework']
    fields = ['uuid', 'Name' ]

    #def get_object(self, queryset=None):
        #obj = User.objects.filter(uuid=self.kwargs['uuid']).first()
        #myAssessment = Assessment.objects.get(pk=self.kwargs['uuid'],analyst=self.kwargs['analyst'])
    #    myAssessment = Assessment.objects.get(analyst=self.kwargs['analyst'])
        #print(myUser.first_name + " " + myUser.last_name)
        #groups = myUser.groups.all()
        #for group in groups:
        #    print(group)
        #print(groups)
        #print("user belongs " + str(groups.count()) + " groups.")
        #for group in groups:
    #        print("retrieved group" + str(group))
    #    return myAssessment

    def get_queryset(self):
        #return Assessment.objects.filter(analyst=self.kwargs['analyst'])
        #print(self.request.user)
        if self.request.user.is_superuser:
            object = Assessment.objects.all()
        else:
            object = Assessment.objects.filter(Q(vendor_contact=self.request.user) | Q(analyst=self.request.user) | Q(owner=self.request.user))
        return object

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        #controls = self.object.frameworkcontrols_set.filter(id=self.kwargs['id']).first()
        #context['controls'] = controls
        #user = self.request.user
        token = Token.objects.get(user=self.request.user)
        context['token'] = token
        return context
